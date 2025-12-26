from docx import Document
import re

def remove_footnotes_and_save(input_docx, output_txt):
    """
    读取Word文件，去除所有脚注内容后保存为TXT
    
    参数:
        input_docx: 输入的Word文件路径(.docx)
        output_txt: 输出的TXT文件路径
    """
    try:
        # 读取Word文档
        doc = Document(input_docx)
        
        # 提取正文内容（不含脚注）
        clean_text = []
        for paragraph in doc.paragraphs:
            clean_text.append(paragraph.text)
        
        # 处理表格中的文本（可选）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    clean_text.append(cell.text)
        
        # 合并文本并去除多余空行
        full_text = "\n".join(clean_text)
        full_text = re.sub(r'\n\s*\n', '\n\n', full_text)  # 保留段落间空行但不多于一个
        
        # 保存为TXT
        with open(output_txt, 'w', encoding='utf-8') as f:
            f.write(full_text)
            
        print(f"成功处理并保存: {output_txt}")
        return True
    
    except Exception as e:
        print(f"处理失败: {str(e)}")
        return False

# 使用示例
if __name__ == "__main__":
    input_file = "input.docx"  # 替换为你的Word文件路径
    output_file = "output.txt"  # 输出的TXT文件路径
    
    remove_footnotes_and_save(input_file, output_file)